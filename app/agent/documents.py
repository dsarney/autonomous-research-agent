from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.models import Source, UploadedDocument

DEFAULT_DOCUMENT_QUERY = (
    "Investigate the uploaded documents: extract the main claims, "
    "assess supporting and contradicting evidence, identify related findings, "
    "and note open questions."
)
TRUNCATION_MARKER = "\n\n[Truncated; remaining text omitted.]"
SNIPPET_CHARS = 400
SEARCH_CONTEXT_CHARS = 1200

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
EXTENSION_CONTENT_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".md": "text/markdown",
}


class DocumentError(ValueError):
    """User-facing validation error for uploaded files."""


@dataclass(frozen=True)
class IncomingFile:
    filename: str
    content_type: str
    data: bytes


def extract_documents(
    files: list[IncomingFile],
    *,
    max_files: int,
    max_upload_mb: int,
    max_chars_per_file: int,
    max_total_chars: int,
) -> list[UploadedDocument]:
    if not files:
        return []
    if len(files) > max_files:
        raise DocumentError(f"You can upload at most {max_files} documents.")

    max_bytes = max_upload_mb * 1024 * 1024
    extracted: list[UploadedDocument] = []
    remaining_total = max_total_chars

    for index, incoming in enumerate(files, start=1):
        filename = sanitize_filename(incoming.filename)
        extension = Path(filename).suffix.lower()
        if extension not in ALLOWED_EXTENSIONS:
            raise DocumentError(
                f"Unsupported file type: {filename}. Use PDF, Word (.docx), text, or Markdown."
            )
        if len(incoming.data) > max_bytes:
            raise DocumentError(f"{filename} is larger than {max_upload_mb} MB.")
        if remaining_total <= 0:
            raise DocumentError(
                "Uploaded documents exceed the total text limit. Remove a file or use shorter papers."
            )

        text = _extract_text(incoming.data, extension, filename)
        excerpt = _truncate(text, min(max_chars_per_file, remaining_total))
        remaining_total = max(0, remaining_total - len(excerpt))
        content_type = incoming.content_type.strip() or EXTENSION_CONTENT_TYPES.get(
            extension, "application/octet-stream"
        )
        extracted.append(
            UploadedDocument(
                id=f"D{index}",
                filename=filename,
                content_type=content_type,
                char_count=len(text),
                excerpt=excerpt,
            )
        )
    return extracted


def documents_to_sources(documents: list[UploadedDocument]) -> list[Source]:
    sources: list[Source] = []
    for document in documents:
        snippet = " ".join(document.excerpt.split())[:SNIPPET_CHARS]
        sources.append(
            Source(
                id=document.id,
                url=f"upload://{document.filename}",
                title=document.filename,
                snippet=snippet,
                relevance=1.0,
                kind="upload",
                credibility_notes="User-uploaded document",
            )
        )
    return sources


def format_document_context(
    documents: list[UploadedDocument] | None,
    *,
    max_chars_each: int | None = None,
) -> str:
    if not documents:
        return ""
    blocks = []
    for document in documents:
        body = document.excerpt
        if max_chars_each is not None:
            body = _truncate(body, max_chars_each)
        blocks.append(f"### {document.id} {document.filename}\n{body}")
    return "Uploaded documents (primary sources):\n\n" + "\n\n".join(blocks)


def sanitize_filename(name: str) -> str:
    base = Path(name or "").name.strip().replace("\x00", "")
    if not base or base in {".", ".."}:
        return "document"
    return base


def _extract_text(data: bytes, extension: str, filename: str) -> str:
    if not data:
        raise DocumentError(f"{filename} is empty.")
    try:
        if extension == ".pdf":
            text = _pdf_text(data)
        elif extension == ".docx":
            text = _docx_text(data)
        else:
            text = data.decode("utf-8", errors="replace")
    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentError(f"{filename} could not be read.") from exc
    cleaned = _normalize_text(text)
    if not cleaned:
        raise DocumentError(
            f"{filename} has no extractable text. Scanned PDFs are not supported."
        )
    return cleaned


def _pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _docx_text(data: bytes) -> str:
    document = Document(BytesIO(data))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _truncate(text: str, limit: int) -> str:
    if limit <= 0:
        return TRUNCATION_MARKER.strip()
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + TRUNCATION_MARKER


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\x00", "").splitlines()]
    collapsed = "\n".join(line for line in lines if line)
    return collapsed.strip()
